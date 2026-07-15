"""
AI Chatbot API Router — The conversational interface for the AI Hiring Chatbot.

The chatbot guides HR/Recruiters through hiring request collection,
then triggers the full agentic workflow (agentic-workflow.md).

Endpoints:
- POST /chatbot/start               — Start a new chatbot session
- POST /chatbot/message             — Send a message and get AI response
- GET  /chatbot/session/{id}        — Get session history
- POST /chatbot/approve-jd          — Approve or reject the generated JD
"""
import uuid
import json
import re
from datetime import datetime
from typing import Optional
from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.messages import SystemMessage, HumanMessage

from backend.config import settings

router = APIRouter(prefix="/chatbot", tags=["AI Chatbot"])

# In-memory session store (in production, use Redis)
_sessions: dict = {}

# Initialize LLM for chatbot
llm = ChatGoogleGenerativeAI(
    model=settings.GEMINI_MODEL,
    temperature=0.3,
    google_api_key=settings.GOOGLE_API_KEY
)

# ── Pydantic Models ────────────────────────────────────────────────────────────

class StartSessionResponse(BaseModel):
    session_id: str
    welcome_message: str
    step: str


class ChatMessageRequest(BaseModel):
    session_id: str
    message: str


class ChatMessageResponse(BaseModel):
    session_id: str
    bot_message: str
    step: str
    hiring_request: Optional[dict] = None
    workflow_triggered: bool = False
    workflow_session_id: Optional[str] = None
    jd_content: Optional[str] = None


class JDApprovalRequest(BaseModel):
    session_id: str
    approved: bool
    feedback: Optional[str] = None


# ── Required fields for a complete hiring request ─────────────────────────────

REQUIRED_FIELDS = {
    "job_title":               "Job Title / Role",
    "skills_required":         "Key Skills & Technologies",
    "candidates_needed":       "Number of Positions",
    "experience_years":        "Experience Required",
    "location":                "Location / Remote Policy",
    "budget":                  "Salary Budget / Range",
}


# ── Regex-based extraction (reliable fallback) ────────────────────────────────

def _regex_extract(message: str) -> dict:
    """
    Extract hiring fields using regex patterns.
    This is the primary reliable extractor — no LLM hallucination possible.
    """
    result = {}
    msg = message.strip()
    msg_lower = msg.lower()

    # ── Job Title / Role ──────────────────────────────────────────────────
    role_patterns = [
        r'(?:role|job\s*title|position|hiring\s*for)\s*[:\-–—]\s*(.+?)(?:\n|,\s*(?:number|experience|skills|key|location|salary|budget|additional)|$)',
        r'^(.+?)\s*(?:developer|engineer|architect|designer|manager|analyst|lead|specialist|consultant|scientist|admin)\b',
    ]
    for pat in role_patterns:
        m = re.search(pat, msg, re.IGNORECASE)
        if m:
            title = m.group(1).strip().rstrip('.,;')
            # If the match is the whole role phrase (e.g. "senior software developer")
            if 'developer' in msg_lower or 'engineer' in msg_lower or 'architect' in msg_lower or 'designer' in msg_lower or 'manager' in msg_lower or 'analyst' in msg_lower or 'lead' in msg_lower or 'scientist' in msg_lower:
                # Try to get the full title including the suffix
                full_match = re.search(
                    r'(?:role|job\s*title|position|hiring\s*for)\s*[:\-–—]\s*(.+?(?:developer|engineer|architect|designer|manager|analyst|lead|specialist|consultant|scientist|admin)\w*)',
                    msg, re.IGNORECASE
                )
                if full_match:
                    title = full_match.group(1).strip().rstrip('.,;')
                else:
                    # Maybe the whole message starts with the role
                    full_match2 = re.search(
                        r'((?:senior|junior|lead|staff|principal|mid|sr|jr)?\s*\w+\s+(?:developer|engineer|architect|designer|manager|analyst|lead|specialist|consultant|scientist))',
                        msg, re.IGNORECASE
                    )
                    if full_match2:
                        title = full_match2.group(1).strip()
            if len(title) > 2:
                result["job_title"] = title.title()
            break

    # ── Number of Positions / Candidates ──────────────────────────────────
    candidates_patterns = [
        r'(?:number\s*of\s*(?:positions?|candidates?|openings?|hires?)|positions?\s*(?:needed|required|to\s*(?:fill|hire)))\s*[:\-–—]?\s*(\d+)',
        r'(?:need|hire|looking\s*for|want)\s+(\d+)\s+(?:candidates?|positions?|people|engineers?|developers?)',
        r'(\d+)\s+(?:positions?|openings?|candidates?|hires?)',
    ]
    for pat in candidates_patterns:
        m = re.search(pat, msg, re.IGNORECASE)
        if m:
            result["candidates_needed"] = int(m.group(1))
            break

    # ── Experience ────────────────────────────────────────────────────────
    exp_patterns = [
        r'(?:experience\s*(?:required|needed|level)?|exp)\s*[:\-–—]?\s*([\d]+[\s\-–—to]+[\d]+\s*(?:years?|yrs?|yr))',
        r'(?:experience\s*(?:required|needed|level)?|exp)\s*[:\-–—]?\s*([\d]+\+?\s*(?:years?|yrs?|yr))',
        r'([\d]+[\s\-–—to]+[\d]+\s*(?:years?|yrs?))\s*(?:of\s*)?(?:experience|exp)',
        r'([\d]+\+?\s*(?:years?|yrs?))\s*(?:of\s*)?(?:experience|exp)',
    ]
    for pat in exp_patterns:
        m = re.search(pat, msg, re.IGNORECASE)
        if m:
            result["experience_years"] = m.group(1).strip()
            break

    # ── Skills / Technologies ─────────────────────────────────────────────
    skills_patterns = [
        r'(?:key\s*)?(?:skills?|tech(?:nologies|nical)?|stack|rechnologies)\s*(?:required|needed|and\s*(?:tech|rech)\w*)?\s*[:\-–—]\s*(.+?)(?:\n|,\s*(?:location|salary|budget|additional|experience|number)|\.?\s*$)',
        r'(?:skills?|tech(?:nologies)?)\s*[:\-–—]\s*(.+?)(?:\.|$)',
    ]
    for pat in skills_patterns:
        m = re.search(pat, msg, re.IGNORECASE)
        if m:
            raw = m.group(1).strip().rstrip('.,;')
            # Split on common delimiters
            skills = [s.strip().rstrip('.,;') for s in re.split(r'[,/\|]|\band\b', raw) if s.strip()]
            # Filter out non-skill items and clean
            cleaned_skills = []
            for s in skills:
                s = s.strip()
                if len(s) > 1 and len(s) < 50:
                    cleaned_skills.append(s)
            if cleaned_skills:
                result["skills_required"] = cleaned_skills
            break

    # ── Location ──────────────────────────────────────────────────────────
    loc_patterns = [
        r'(?:location|city|place|work\s*(?:location|from)|based\s*(?:in|at))\s*[:\-–—]\s*(.+?)(?:\n|,\s*(?:salary|budget|additional)|\.?\s*$)',
        r'(?:location|city)\s*[:\-–—]\s*(.+?)(?:\.|,|$)',
    ]
    for pat in loc_patterns:
        m = re.search(pat, msg, re.IGNORECASE)
        if m:
            loc = m.group(1).strip().rstrip('.,;')
            if len(loc) > 1:
                result["location"] = loc.title()
            break

    # ── Budget / Salary ───────────────────────────────────────────────────
    budget_patterns = [
        r'(?:salary|budget|compensation|pay|ctc|package|sal)\s*(?:range|budget)?\s*[:\-–—]\s*(.+?)(?:\n|,\s*(?:additional|location)|\.?\s*(?:go\s|start\s|$))',
        r'((?:₹|rs\.?|inr|usd|\$)\s*[\d]+[\s\-–—to]+[\d]+\s*(?:lpa|lakhs?|lacs?|k|cr|crore)?)',
        r'([\d]+\s*(?:to|-)\s*[\d]+\s*(?:lpa|lakhs?|lacs?|ctc|k))',
    ]
    for pat in budget_patterns:
        m = re.search(pat, msg, re.IGNORECASE)
        if m:
            budget = m.group(1).strip().rstrip('.,;')
            if len(budget) > 1:
                result["budget"] = budget
            break

    # ── Additional Requirements ───────────────────────────────────────────
    add_patterns = [
        r'(?:additional\s*(?:requirements?|info|details?|notes?))\s*[:\-–—]\s*(.+?)(?:\n|$)',
    ]
    for pat in add_patterns:
        m = re.search(pat, msg, re.IGNORECASE)
        if m:
            req = m.group(1).strip().rstrip('.,;')
            if len(req) > 2:
                result["additional_requirements"] = req
            break

    return result


# ── LLM-based extraction (used as secondary approach) ─────────────────────────

EXTRACT_PROMPT = """You are a strict data extraction assistant. Extract hiring-related fields from this message.

STRICT RULES:
1. ONLY extract values the user EXPLICITLY wrote. NEVER guess or infer.
2. If a field is NOT mentioned, DO NOT include it.
3. Return ONLY a JSON object.

Fields to look for:
- "job_title": string — the job role (e.g. "Senior Python Developer")
- "skills_required": array of strings — technologies/skills (e.g. ["Python", "FastAPI", "AWS"])
- "candidates_needed": integer — number of positions (e.g. 3)
- "experience_years": string — experience needed (e.g. "5 years")
- "location": string — work location (e.g. "Bengaluru")
- "budget": string — salary range (e.g. "18-20 LPA CTC")
- "additional_requirements": string — extra requirements

User's message:
"{message}"

Return ONLY valid JSON. No explanation."""


def _llm_extract(message: str) -> dict:
    """Use LLM to extract fields. Secondary to regex extraction."""
    prompt = EXTRACT_PROMPT.format(message=message)
    try:
        response = llm.invoke(prompt)
        content = response.content.strip()
        json_match = re.search(r'\{.*\}', content, re.DOTALL)
        if json_match:
            extracted = json.loads(json_match.group())
            cleaned = {}
            for k, v in extracted.items():
                if v is None:
                    continue
                if isinstance(v, str) and not v.strip():
                    continue
                if isinstance(v, list) and len(v) == 0:
                    continue
                cleaned[k] = v
            return cleaned
    except Exception as e:
        print(f"[Chatbot] LLM extraction error: {e}")
    return {}


def _extract_all_fields(message: str, existing_data: dict) -> dict:
    """
    Extract all hiring fields from a user message.
    Uses regex first (reliable, no hallucination), then fills gaps with LLM.
    """
    # Step 1: Regex extraction (primary — deterministic, no hallucination)
    regex_result = _regex_extract(message)
    print(f"[Chatbot] Regex extracted: {regex_result}")

    # Step 2: LLM extraction (secondary — fills gaps regex missed)
    llm_result = _llm_extract(message)
    print(f"[Chatbot] LLM extracted: {llm_result}")

    # Merge: regex takes priority, LLM fills gaps
    merged = {}
    all_keys = set(list(regex_result.keys()) + list(llm_result.keys()))
    for key in all_keys:
        if key in regex_result:
            merged[key] = regex_result[key]
        elif key in llm_result:
            merged[key] = llm_result[key]

    # Don't overwrite existing data with new extraction unless it's a new field
    final = {}
    for key, val in merged.items():
        if key not in existing_data or not existing_data.get(key):
            final[key] = val
        elif key in existing_data:
            # Only overwrite if the user explicitly mentioned it again
            # (we assume they did if it was extracted)
            final[key] = val

    print(f"[Chatbot] Final merged extraction: {final}")
    return final


# ── Field checking ─────────────────────────────────────────────────────────────

def _get_missing_required_fields(hiring_request: dict) -> dict:
    """Return dict of {field_key: human_label} for required fields still missing."""
    missing = {}
    for field_key, label in REQUIRED_FIELDS.items():
        val = hiring_request.get(field_key)
        if val is None:
            missing[field_key] = label
        elif isinstance(val, str) and not val.strip():
            missing[field_key] = label
        elif isinstance(val, list) and len(val) == 0:
            missing[field_key] = label
    return missing


def _user_wants_to_proceed(message: str) -> bool:
    """Check if user is signaling they want to skip remaining fields and proceed."""
    proceed_phrases = [
        "go ahead", "start build", "build the jd", "generate jd", "generate the jd",
        "create jd", "create the jd", "proceed", "that's it", "thats it",
        "that's all", "thats all", "skip", "no more", "nothing else",
        "none", "na", "nah", "n/a", "no additional", "no other",
        "start the workflow", "kick off", "let's go", "lets go",
    ]
    msg_lower = message.lower().strip()
    return any(phrase in msg_lower for phrase in proceed_phrases)


# ── Summary builder ────────────────────────────────────────────────────────────

def _build_summary(hr: dict) -> str:
    """Build a markdown summary of the hiring request."""
    skills = hr.get("skills_required", [])
    skills_str = ", ".join(skills) if isinstance(skills, list) else str(skills) if skills else "Not specified"
    return (
        f"📋 **Hiring Request Summary**\n\n"
        f"- **Job Title**: {hr.get('job_title', 'Not specified')}\n"
        f"- **Skills & Technologies**: {skills_str}\n"
        f"- **Experience**: {hr.get('experience_years', 'Not specified')}\n"
        f"- **Location**: {hr.get('location', 'Not specified')}\n"
        f"- **Budget/Salary**: {hr.get('budget', 'Not specified')}\n"
        f"- **Candidates to Hire**: {hr.get('candidates_needed', 'Not specified')}\n"
        f"- **Additional Requirements**: {hr.get('additional_requirements', 'None')}"
    )


# ── JD generation ─────────────────────────────────────────────────────────────

async def _generate_jd_in_session(session: dict) -> str:
    """Generate a JD based on the hiring request and return it formatted for chat."""
    hr = session.get("hiring_request", {})
    skills = hr.get("skills_required", [])
    skills_str = ", ".join(skills) if isinstance(skills, list) else str(skills) if skills else "Not specified"
    feedback = hr.get("jd_feedback", "")

    # Detect if user explicitly asked for more detail / a longer JD
    additional = hr.get("additional_requirements", "") or ""
    user_wants_long = any(kw in additional.lower() for kw in [
        "detailed", "elaborate", "comprehensive", "long", "full", "extensive",
        "in-depth", "include everything", "all sections", "more sections",
    ])

    size_note = (
        "The user explicitly requested a detailed / comprehensive JD — include all sections "
        "and expand each one with thorough, specific content."
        if user_wants_long else
        "Keep the JD CONCISE. Use at most 5 bullet points per section. "
        "Do NOT pad or add generic filler. Every bullet must be specific to this exact role and stack."
    )

    jd_prompt = f"""You are a senior technical recruiter writing a real, professional Job Description.

## Hiring Details
- **Role:** {hr.get('job_title', 'Software Engineer')}
- **Required Skills:** {skills_str}
- **Experience:** {hr.get('experience_years', '3+ years')}
- **Location:** {hr.get('location', 'Remote')}
- **Salary / Budget:** {hr.get('budget', 'Competitive')}
- **Openings:** {hr.get('candidates_needed', 1)}
{f"- **Additional context:** {additional}" if additional and additional.strip().lower() not in ['none', 'na', 'n/a', ''] else ""}
{f"- **Revise based on this feedback:** {feedback}" if feedback else ""}

## Output Rules — FOLLOW STRICTLY
1. Write ONLY content directly relevant to this specific role and the listed tech stack.
2. NEVER use generic buzzwords: no "fast-paced", "passion for excellence", "wear many hats", "rockstar", "ninja", "synergy", "dynamic team", "go-getter".
3. Do NOT invent or assume company perks, culture, or benefits that were not provided. Omit a "What We Offer" section entirely unless the user gave specific benefit details.
4. Do NOT include a "How to Apply" section unless the user explicitly asked for one.
5. Every responsibility and skill bullet must be concrete, actionable, and specific to {hr.get('job_title', 'the role')} using {skills_str}.
6. {size_note}
7. Use clean Markdown: `##` for headings, `**label:**` for inline labels, `-` for bullet points.
8. Start directly with the JD — no preamble like "Here is the JD:", "Sure!", or "Certainly!".
9. The top of the JD must show: Role, Location, Experience, Salary, and Openings as a compact info block.

## Sections to Include (in this order)
1. **[Role Title]** — one-liner tagline describing the role
2. **Overview** — 2–3 sentences: what the person owns, who they work with, impact of the role
3. **Key Responsibilities** — concrete duties using the specified stack
4. **Required Skills & Experience** — only the explicitly provided skills and experience levels
5. **Good to Have** — optional, ≤3 items, only if genuinely relevant to the stack; skip if nothing fits

Write the Job Description now:
"""

    try:
        response = llm.invoke(jd_prompt)
        jd_content = response.content.strip()
        session["jd_content"] = jd_content

        return (
            f"✨ **Here's your Job Description:**\n\n"
            f"{jd_content}\n\n"
            f"---\n"
            f"👆 Please review the JD above. "
            f"Type **'approve'** to post it on all platforms and start the hiring workflow, "
            f"or tell me what changes you'd like to make."
        )
    except Exception as e:
        return f"Sorry, I encountered an error generating the JD: {str(e)}. Please try again."


# ── Workflow trigger ──────────────────────────────────────────────────────────

async def _trigger_hiring_workflow(session: dict) -> str:
    """Trigger the LangGraph hiring workflow with the collected hiring request."""
    import uuid as _uuid
    workflow_id = str(_uuid.uuid4())[:8].upper()

    hr = session.get("hiring_request", {})
    job_title = hr.get("job_title", "Open Position")

    try:
        from backend.workflows.graph import hiring_graph

        thread_config = {"configurable": {"thread_id": workflow_id}}
        initial_state = {
            "messages": [],
            "job_id": workflow_id,
            "goal": f"Hire {hr.get('candidates_needed', 1)} {job_title}(s)",
            "next_action": "jd_generation",
            "agent_statuses": {"chatbot": "initiated"},
            "plan": [],
            "current_step_index": 0,
            "data": {},
            "hiring_request": hr,
            "jd_content": session.get("jd_content"),
            "jd_approved": True,
            "jd_feedback": "",
            "jd_retry_count": 1,
            "posting_status": {},
            "application_count": 0,
            "candidates_needed": hr.get("candidates_needed", 10),
            "sourcing_retry_count": 0,
            "shortlisted_candidates": [],
            "candidate_rankings": {},
            "scheduled_interviews": [],
            "interview_results": {},
            "selected_candidates": [],
            "offer_letters": {},
            "offer_status": {},
            "negotiation_rounds": {},
            "onboarding_tasks": [],
            "onboarding_status": {},
            "chat_session_id": session.get("session_id"),
            "chat_history": session.get("messages", []),
            "error": None,
        }

        import asyncio
        asyncio.create_task(
            asyncio.to_thread(
                hiring_graph.invoke,
                initial_state,
                thread_config
            )
        )
        print(f"[Chatbot] ✅ Workflow {workflow_id} triggered for job: {job_title}")
    except Exception as e:
        print(f"[Chatbot] Workflow trigger error: {e}")

    # ── Persist a Job record so it appears in Job Management ─────────────────
    try:
        from backend.database.session import AsyncSessionLocal
        from backend.database.models import Job, JobDescription, User
        from sqlalchemy import select as _select

        # Map human-readable job type strings to DB enum values
        _job_type_map = {
            "full time": "full_time", "fulltime": "full_time",
            "part time": "part_time", "parttime": "part_time",
            "contract": "contract", "freelance": "contract",
            "remote": "remote",
        }
        raw_job_type = str(hr.get("job_type", "")).lower().strip()
        job_type_val = _job_type_map.get(raw_job_type, "full_time")

        async with AsyncSessionLocal() as db:
            # Use the first active user as creator (chatbot has no auth context)
            user_result = await db.execute(
                _select(User).where(User.is_active == True).limit(1)
            )
            first_user = user_result.scalar_one_or_none()
            creator_id = first_user.id if first_user else "chatbot"

            job = Job(
                title=job_title,
                department=str(hr.get("department", "General")),
                location=str(hr.get("location", "Not specified")),
                job_type=job_type_val,
                experience_level=str(hr.get("experience_years", "Not specified")),
                hiring_goal=f"Hire {hr.get('candidates_needed', 1)} {job_title}(s)",
                target_candidate_count=int(hr.get("candidates_needed", 1)),
                status="approved",
                created_by=creator_id,
            )
            db.add(job)
            await db.flush()  # get job.id before commit

            if session.get("jd_content"):
                jd = JobDescription(
                    job_id=job.id,
                    content=session["jd_content"],
                    version=1,
                    status="approved",
                )
                db.add(jd)

            # Store the real DB job_id back into the session for reference
            session["db_job_id"] = job.id
            await db.commit()
            print(f"[Chatbot] ✅ Job record created in DB: {job.id} ({job_title})")
    except Exception as e:
        print(f"[Chatbot] ⚠️  Failed to persist Job to DB: {e}")

    return workflow_id


# ── API Endpoints ──────────────────────────────────────────────────────────────

@router.post("/start", response_model=StartSessionResponse)
async def start_chatbot_session():
    """Start a new AI Hiring Chatbot session."""
    session_id = str(uuid.uuid4())

    welcome_message = (
        "👋 Hi there! I'm your **AI Hiring Assistant**.\n\n"
        "To get started, tell me about the position you'd like to fill. "
        "Please include as many of these details as you can:\n\n"
        "- **Role** (e.g. Senior Python Developer)\n"
        "- **Number of Positions**\n"
        "- **Experience Required** (e.g. 3-5 years)\n"
        "- **Key Skills & Technologies** (e.g. Python, FastAPI, AWS)\n"
        "- **Location** (e.g. Remote, Bangalore)\n"
        "- **Salary Budget** (e.g. ₹15-25 LPA)\n"
        "- **Additional Requirements** (optional)\n\n"
        "You can provide everything in one message or we can go step by step — your choice! 🚀"
    )

    _sessions[session_id] = {
        "session_id": session_id,
        "step": "collect_details",
        "hiring_request": {},
        "messages": [
            {"role": "assistant", "content": welcome_message, "timestamp": datetime.utcnow().isoformat()}
        ],
        "jd_content": None,
        "workflow_session_id": None,
        "created_at": datetime.utcnow().isoformat(),
    }

    return StartSessionResponse(
        session_id=session_id,
        welcome_message=welcome_message,
        step="collect_details"
    )


@router.post("/message", response_model=ChatMessageResponse)
async def send_message(request: ChatMessageRequest):
    """
    Send a message to the chatbot and get an AI response.
    Smart extraction: regex first, LLM second. Never asks for the same info twice.
    """
    session = _sessions.get(request.session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found. Please start a new session.")

    # Add user message to history
    session["messages"].append({
        "role": "user",
        "content": request.message,
        "timestamp": datetime.utcnow().isoformat()
    })

    current_step = session.get("step", "collect_details")
    workflow_triggered = False
    workflow_session_id = None

    # ── Handle simple greetings at the very start (no data yet) ───────────────
    msg_clean = request.message.lower().strip().replace(".", "").replace("!", "")
    if current_step == "collect_details" and not session.get("hiring_request") and msg_clean in [
        "hi", "hello", "hey", "greetings", "hi there", "hello there",
    ]:
        bot_message = (
            "Great! Let's set up your hiring request. 🎯\n\n"
            "Please provide the following details:\n\n"
            "- **Role** (e.g. Senior Python Developer)\n"
            "- **Number of Positions**\n"
            "- **Experience Required**\n"
            "- **Key Skills & Technologies**\n"
            "- **Location**\n"
            "- **Salary Budget**\n"
            "- **Additional Requirements** (optional)\n\n"
            "You can share everything in one message!"
        )
        session["messages"].append({
            "role": "assistant",
            "content": bot_message,
            "timestamp": datetime.utcnow().isoformat()
        })
        return ChatMessageResponse(
            session_id=request.session_id,
            bot_message=bot_message,
            step=current_step,
            hiring_request=session.get("hiring_request"),
            workflow_triggered=False,
            workflow_session_id=None
        )

    # ── Data collection phase ─────────────────────────────────────────────────
    if current_step == "collect_details":
        # Extract ALL fields from the user's message (regex + LLM)
        extracted = _extract_all_fields(request.message, session.get("hiring_request", {}))
        if extracted:
            session["hiring_request"].update(extracted)

        # Check what's still missing
        missing = _get_missing_required_fields(session["hiring_request"])
        user_wants_proceed = _user_wants_to_proceed(request.message)

        if len(missing) == 0:
            # All required fields collected → generate JD immediately
            session["step"] = "jd_generation"
            bot_message = (
                "✅ **Hiring request details completed!**\n\n"
                "I'm now generating a professional Job Description tailored to your requirements. "
                "This will take just a moment... 🤖⚙️"
            )
            session["messages"].append({
                "role": "assistant",
                "content": bot_message,
                "timestamp": datetime.utcnow().isoformat()
            })
            bot_message = await _generate_jd_in_session(session)
            session["step"] = "jd_review"
        elif user_wants_proceed and session.get("hiring_request", {}).get("job_title"):
            # User wants to proceed even with missing fields — fill defaults and generate JD immediately
            hr = session["hiring_request"]
            defaults = {
                "candidates_needed": 1,
                "experience_years": "Not specified",
                "location": "Not specified",
                "budget": "Competitive",
                "skills_required": [],
            }
            for key, default_val in defaults.items():
                if key not in hr or not hr[key]:
                    hr[key] = default_val

            session["step"] = "jd_generation"
            bot_message = (
                "✅ **Proceeding to generate Job Description with available details...** 🤖⚙️"
            )
            session["messages"].append({
                "role": "assistant",
                "content": bot_message,
                "timestamp": datetime.utcnow().isoformat()
            })
            bot_message = await _generate_jd_in_session(session)
            session["step"] = "jd_review"
        else:
            # Still missing some fields — ask for them
            missing_labels = list(missing.values())
            missing_list = "\n".join(f"- **{label}**" for label in missing_labels)

            if extracted:
                # Acknowledged what was provided, ask for the rest
                collected_keys = [k for k in extracted.keys()]
                collected_names = [REQUIRED_FIELDS.get(k, k) for k in collected_keys if k in REQUIRED_FIELDS]
                ack = ""
                if collected_names:
                    ack = f"Thanks! I've noted the **{', '.join(collected_names)}**. "

                bot_message = (
                    f"{ack}I still need a few more details:\n\n"
                    f"{missing_list}\n\n"
                    "Please provide these, or say **'go ahead'** to proceed with what we have."
                )
            else:
                # Nothing extracted — but don't say "I couldn't catch" if they just typed something short
                bot_message = (
                    "I need the following details to set up your hiring request:\n\n"
                    f"{missing_list}\n\n"
                    "Please share these details, or say **'go ahead'** to proceed with what we have."
                )

    # ── Confirmation step ─────────────────────────────────────────────────────
    elif current_step == "confirmation":
        msg_lower = request.message.lower()
        if any(word in msg_lower for word in ["yes", "confirm", "looks good", "correct", "proceed", "go ahead", "ok", "okay", "lgtm"]):
            session["step"] = "jd_generation"
            bot_message = (
                "✅ **Hiring request confirmed!**\n\n"
                "I'm now generating a professional Job Description tailored to your requirements. "
                "This will take just a moment... 🤖⚙️"
            )
            session["messages"].append({
                "role": "assistant",
                "content": bot_message,
                "timestamp": datetime.utcnow().isoformat()
            })

            # Generate JD
            bot_message = await _generate_jd_in_session(session)
            session["step"] = "jd_review"
        else:
            # User wants to edit — extract any updates
            extracted = _extract_all_fields(request.message, session.get("hiring_request", {}))
            if extracted:
                session["hiring_request"].update(extracted)

            summary = _build_summary(session["hiring_request"])
            bot_message = (
                f"Got it! I've updated the details. Here's the revised summary:\n\n"
                f"{summary}\n\n"
                "Does everything look correct now? Type **'confirm'** to proceed."
            )

    # ── JD Review step ────────────────────────────────────────────────────────
    elif current_step == "jd_review":
        msg_lower = request.message.lower()
        if any(word in msg_lower for word in ["approve", "approved", "looks good", "yes", "go ahead", "post it", "publish"]):
            # JD approved → trigger hiring workflow immediately
            session["step"] = "workflow_running"
            workflow_triggered = True
            workflow_session_id = await _trigger_hiring_workflow(session)
            session["workflow_session_id"] = workflow_session_id
            bot_message = (
                "🚀 **JD Approved! Hiring Workflow is now running!**\n\n"
                "The complete pipeline has been kicked off. Candidates will be sourced "
                "from the **Hiring Platform** (http://localhost:8001) where applicants upload their resumes.\n\n"
                "Pipeline stages:\n"
                "1. ✅ Job Description approved\n"
                "2. 👀 Sourcing candidates from the Hiring Platform\n"
                "3. 📋 Auto-shortlisting & ranking candidates\n"
                "4. 📅 Scheduling & conducting interviews\n"
                "5. 📄 Generating offer letters\n"
                "6. 🎊 Completing onboarding\n\n"
                "Track live progress on the **Workflow Monitor** page. 🚀"
            )
            session["step"] = "complete"
        else:
            session["hiring_request"]["jd_feedback"] = request.message
            bot_message = (
                f"Got it! I'll revise the JD based on your feedback: *\"{request.message}\"*\n\n"
                "Regenerating the JD now... 🔄"
            )
            session["messages"].append({
                "role": "assistant",
                "content": bot_message,
                "timestamp": datetime.utcnow().isoformat()
            })
            bot_message = await _generate_jd_in_session(session)

    else:
        bot_message = "I'm not sure how to help with that right now. Please start a new session if needed."

    # Add bot response to history
    session["messages"].append({
        "role": "assistant",
        "content": bot_message,
        "timestamp": datetime.utcnow().isoformat()
    })

    return ChatMessageResponse(
        session_id=request.session_id,
        bot_message=bot_message,
        step=session.get("step", current_step),
        hiring_request=session.get("hiring_request"),
        workflow_triggered=workflow_triggered,
        workflow_session_id=workflow_session_id,
        jd_content=session.get("jd_content")
    )


@router.get("/session/{session_id}")
async def get_session(session_id: str):
    """Get chatbot session history."""
    session = _sessions.get(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    return {
        "session_id": session_id,
        "step": session.get("step"),
        "messages": session.get("messages", []),
        "hiring_request": session.get("hiring_request", {}),
        "jd_content": session.get("jd_content"),
        "workflow_session_id": session.get("workflow_session_id"),
    }


@router.post("/approve-jd")
async def approve_jd(request: JDApprovalRequest):
    """Approve or reject the generated JD."""
    session = _sessions.get(request.session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if request.approved:
        session["jd_approved"] = True
        session["step"] = "workflow_running"
        workflow_session_id = await _trigger_hiring_workflow(session)
        session["workflow_session_id"] = workflow_session_id
        return {
            "approved": True,
            "message": "JD approved. Hiring workflow started!",
            "workflow_session_id": workflow_session_id
        }
    else:
        session["jd_approved"] = False
        session["hiring_request"]["jd_feedback"] = request.feedback or "Please improve the JD"
        session["step"] = "jd_generation"
        return {
            "approved": False,
            "message": "JD sent back for revision.",
            "feedback": request.feedback
        }


# ── JD Download Endpoints ──────────────────────────────────────────────────────

@router.get("/session/{session_id}/download-jd/pdf")
async def download_jd_pdf(session_id: str):
    """Download the generated JD as a PDF file."""
    session = _sessions.get(session_id)
    if not session or not session.get("jd_content"):
        raise HTTPException(status_code=404, detail="Session or Job Description not found")

    import io
    from fastapi.responses import StreamingResponse
    
    title = f"Job Description - {session['hiring_request'].get('job_title', 'Open Position')}"
    text = session["jd_content"]
    
    # Generate PDF using PyMuPDF (fitz)
    try:
        import fitz

        def _strip_inline_bold(s: str) -> str:
            """Remove **...** markers and return plain text (used for width estimation)."""
            return re.sub(r'\*\*(.+?)\*\*', r'\1', s)

        def _parse_bold_segments(s: str):
            """
            Split a string into segments: [(text, is_bold), ...]
            e.g. '**Foo:** bar' -> [('Foo:', True), (' bar', False)]
            """
            segments = []
            pattern = re.compile(r'\*\*(.+?)\*\*')
            last = 0
            for m in pattern.finditer(s):
                if m.start() > last:
                    segments.append((s[last:m.start()], False))
                segments.append((m.group(1), True))
                last = m.end()
            if last < len(s):
                segments.append((s[last:], False))
            return segments if segments else [(s, False)]

        def _insert_mixed_line(page, x_start, y, line_text, base_fontsize, base_bold, color, max_width):
            """
            Insert a line that may contain inline **bold** segments.
            Returns the new y position after inserting.
            """
            segments = _parse_bold_segments(line_text)
            x = x_start
            # Estimate character width (approx)
            char_w = base_fontsize * 0.5
            for seg_text, seg_bold in segments:
                fn = "helvetica-bold" if (base_bold or seg_bold) else "helvetica"
                # Check if we need to wrap — simple approach: write word by word
                words = seg_text.split(' ')
                for i, word in enumerate(words):
                    token = word if i == len(words) - 1 else word + ' '
                    token_w = len(token) * char_w
                    if x + token_w > x_start + max_width and x > x_start:
                        x = x_start
                        y += base_fontsize + 4
                    if token.strip():
                        page.insert_text((x, y), token, fontsize=base_fontsize, fontname=fn, color=color)
                    x += token_w
            return y + base_fontsize + 4

        doc = fitz.open()
        page = doc.new_page()
        margin = 50
        width = page.rect.width - (margin * 2)

        y = margin + 20

        # Title
        page.insert_text((margin, y), title, fontsize=16, fontname="helvetica-bold", color=(0.1, 0.1, 0.4))
        y += 30

        for line in text.split("\n"):
            line = line.strip()
            if not line:
                y += 8
                continue

            if y > page.rect.height - margin - 30:
                page = doc.new_page()
                y = margin + 20

            fontsize = 10
            is_heading = False
            color = (0.2, 0.2, 0.2)

            if line.startswith("# "):
                line = line[2:]
                fontsize = 14
                is_heading = True
                color = (0.1, 0.1, 0.4)
                y += 8
            elif line.startswith("## "):
                line = line[3:]
                fontsize = 12
                is_heading = True
                color = (0.15, 0.15, 0.35)
                y += 6
            elif line.startswith("### "):
                line = line[4:]
                fontsize = 11
                is_heading = True
                color = (0.2, 0.2, 0.3)
                y += 4
            elif line.startswith("- ") or line.startswith("* "):
                page.insert_text((margin, y), "•", fontsize=10, fontname="helvetica", color=(0.3, 0.3, 0.3))
                line = line[2:]
                y = _insert_mixed_line(page, margin + 15, y, line, fontsize, False, color, width - 15)
                continue

            y = _insert_mixed_line(page, margin, y, line, fontsize, is_heading, color, width)

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        headers = {
            "Content-Disposition": f"attachment; filename=JD_{session_id[:8]}.pdf",
            "Access-Control-Expose-Headers": "Content-Disposition"
        }
        return StreamingResponse(file_stream, media_type="application/pdf", headers=headers)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")


@router.get("/session/{session_id}/download-jd/doc")
async def download_jd_docx(session_id: str):
    """Download the generated JD as a Word document (.docx)."""
    session = _sessions.get(session_id)
    if not session or not session.get("jd_content"):
        raise HTTPException(status_code=404, detail="Session or Job Description not found")

    import io
    from fastapi.responses import StreamingResponse
    
    title = f"Job Description - {session['hiring_request'].get('job_title', 'Open Position')}"
    text = session["jd_content"]
    
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor

        def _add_paragraph_with_bold(doc_obj, raw_line: str, style=None):
            """
            Add a paragraph to the docx, parsing **bold** inline markers
            so they are rendered as actual bold runs instead of raw asterisks.
            """
            para = doc_obj.add_paragraph(style=style) if style else doc_obj.add_paragraph()
            pattern = re.compile(r'\*\*(.+?)\*\*')
            last = 0
            for m in pattern.finditer(raw_line):
                if m.start() > last:
                    para.add_run(raw_line[last:m.start()])
                bold_run = para.add_run(m.group(1))
                bold_run.bold = True
                last = m.end()
            if last < len(raw_line):
                para.add_run(raw_line[last:])
            return para

        doc = Document()
        doc.add_heading(title, level=1)

        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("# "):
                doc.add_heading(re.sub(r'\*\*(.+?)\*\*', r'\1', line[2:]), level=1)
            elif line.startswith("## "):
                doc.add_heading(re.sub(r'\*\*(.+?)\*\*', r'\1', line[3:]), level=2)
            elif line.startswith("### "):
                doc.add_heading(re.sub(r'\*\*(.+?)\*\*', r'\1', line[4:]), level=3)
            elif line.startswith("- ") or line.startswith("* "):
                _add_paragraph_with_bold(doc, line[2:], style='List Bullet')
            else:
                _add_paragraph_with_bold(doc, line)

        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        headers = {
            "Content-Disposition": f"attachment; filename=JD_{session_id[:8]}.docx",
            "Access-Control-Expose-Headers": "Content-Disposition"
        }
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to generate Word document: {str(e)}")
